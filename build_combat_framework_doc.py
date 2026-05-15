from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\PK")
MD = ROOT / "战斗表现底层框架v0.1.md"
OUT = ROOT / "设计文档" / "战斗表现底层框架v0.1.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Subtitle", 11, "555555"),
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name == "Title" or name.startswith("Heading"):
            style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(40, 40, 40)


def add_table_from_md(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, text in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.paragraphs[0].add_run(text).bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
    set_table_borders(table)
    doc.add_paragraph()


def parse_md_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        parts = [p.strip() for p in raw.strip("|").split("|")]
        if not all(set(p) <= set("-: ") for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def build():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    configure_styles(doc)

    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_md_table(lines, i)
            if rows:
                add_table_from_md(doc, rows)
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(stripped[2:])
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:])
        elif len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:4]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(stripped.split(". ", 1)[1])
        else:
            p = doc.add_paragraph()
            if stripped.startswith("底层规则：") or stripped.startswith("注意：") or stripped.startswith("适用：") or stripped.startswith("必须包含：") or stripped.startswith("避免：") or stripped.startswith("承载：") or stripped.startswith("可用通道：") or stripped.startswith("示例："):
                head, rest = stripped.split("：", 1)
                p.add_run(head + "：").bold = True
                p.add_run(rest)
            else:
                p.add_run(stripped)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
