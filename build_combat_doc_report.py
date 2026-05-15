from pathlib import Path
import math
import cv2
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"E:\PK")
VIDEO = ROOT / "combat_video_analysis" / "combat_spider_cave.mp4"
OUT_DIR = ROOT / "设计文档" / "战斗表现分析报告"
BOARD_DIR = OUT_DIR / "证据截图板"
DOCX_PATH = OUT_DIR / "战斗DEMO表现分析报告-蜘蛛洞穴.docx"


def font(size=28, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()


FONT_TITLE = font(30, True)
FONT_LABEL = font(23, False)
FONT_SMALL = font(19, False)


def extract_frame(cap, t, tile_w=460, tile_h=860, label=""):
    cap.set(cv2.CAP_PROP_POS_MSEC, t * 1000)
    ok, frame = cap.read()
    if not ok:
        raise RuntimeError(f"Failed to read frame at {t}s")
    img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
    img.thumbnail((tile_w, tile_h - 90))
    canvas = Image.new("RGB", (tile_w, tile_h), (18, 20, 22))
    canvas.paste(img, ((tile_w - img.width) // 2, 0))
    draw = ImageDraw.Draw(canvas)
    draw.rectangle([0, tile_h - 90, tile_w, tile_h], fill=(26, 30, 36))
    draw.text((14, tile_h - 78), f"{t:05.1f}s", font=FONT_LABEL, fill=(255, 230, 150))
    draw.text((14, tile_h - 45), label, font=FONT_SMALL, fill=(235, 238, 245))
    return canvas


def make_board(name, title, frames):
    cap = cv2.VideoCapture(str(VIDEO))
    if not cap.isOpened():
        raise RuntimeError(f"Cannot open video: {VIDEO}")
    cols, rows = 3, 2
    tile_w, tile_h = 460, 860
    title_h = 78
    board = Image.new("RGB", (cols * tile_w, rows * tile_h + title_h), (10, 12, 15))
    draw = ImageDraw.Draw(board)
    draw.rectangle([0, 0, board.width, title_h], fill=(18, 42, 62))
    draw.text((24, 20), title, font=FONT_TITLE, fill=(255, 255, 255))
    for idx, (t, label) in enumerate(frames):
        tile = extract_frame(cap, t, tile_w, tile_h, label)
        x = (idx % cols) * tile_w
        y = title_h + (idx // cols) * tile_h
        board.paste(tile, (x, y))
    cap.release()
    path = BOARD_DIR / f"{name}.jpg"
    board.save(path, quality=92)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    return p


def add_issue(doc, num, title, diagnosis, evidence, advice, image_path):
    doc.add_heading(f"{num}. {title}", level=1)
    p = doc.add_paragraph()
    p.add_run("诊断：").bold = True
    p.add_run(diagnosis)
    p = doc.add_paragraph()
    p.add_run("录屏证据：").bold = True
    p.add_run(evidence)
    p = doc.add_paragraph()
    p.add_run("建议：").bold = True
    p.add_run(advice)
    doc.add_picture(str(image_path), width=Inches(6.3))
    add_caption(doc, f"图 {num}：{title} - 关键论证帧合集")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Subtitle", 11, "555555"),
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        if "Heading" in name or name == "Title":
            s.font.bold = True

    cap = styles["Caption"]
    cap.font.name = "Microsoft YaHei"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor(90, 90, 90)


def build_doc(boards):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("战斗 DEMO 表现分析报告")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("蜘蛛洞穴战斗录屏 | 战斗表现问题挖掘、证据帧与迭代方案")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_borders(meta, "E5E7EB", "4")
    rows = [
        ("分析对象", r"E:/PK/设计文档/战斗录屏-蜘蛛洞穴.mp4"),
        ("视频信息", "竖屏 500x1078，30fps，约 10分43秒"),
        ("分析方式", "按整体观感、命中链路、敌人威胁、VFX/SFX/UI、节奏与 Boss 阶段进行分层审片"),
        ("输出目的", "为当前战斗 demo 的表现迭代建立问题证据、改动优先级和候选框架规则"),
    ]
    for r, (k, v) in zip(meta.rows, rows):
        r.cells[0].width = Inches(1.4)
        r.cells[1].width = Inches(4.9)
        set_cell_shading(r.cells[0], "F2F4F7")
        for c in r.cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
        r.cells[0].paragraphs[0].add_run(k).bold = True
        r.cells[1].paragraphs[0].add_run(v)

    doc.add_heading("一、核心结论", level=1)
    p = doc.add_paragraph()
    p.add_run("这套 demo 已经形成了“第一人称魔法 + 洞穴蜘蛛遭遇 + 肉鸽卡牌成长”的基本身份。").bold = True
    p.add_run("当前主要短板不是缺少特效，而是关键信息没有形成清晰层级：玩家经常能看到强烈的蓝白技能、伤害数字和敌人血条变化，但不总能稳定判断“技能打中了谁、敌人受到了什么后果、自己为什么受伤、下一秒需要防什么”。")

    doc.add_paragraph("建议第一轮迭代优先处理三件事：")
    add_numbered(doc, [
        "压低持续性大面积 VFX 的遮挡，让敌人剪影、接触点和血条重新成为画面主信息。",
        "为普通敌人和 Boss 建立分级受击反应，让命中从“数字变化”升级为“身体被影响”。",
        "为蜘蛛网、毒液、Boss 攻击补足前摇和来源提示，让玩家受击变得更公平、更可学习。",
    ])

    doc.add_heading("二、问题论证与证据帧", level=1)
    add_issue(
        doc,
        1,
        "命中反馈被大面积 VFX 淹没",
        "冰系技能的视觉量足够强，但接触点、受击目标和敌人轮廓经常被蓝白冰雾、冰爆和持续光效盖住。玩家能感到技能很强，却不总能确认具体命中关系。",
        "0:24-0:31 的小蜘蛛战斗和 6:34-7:18 的多敌人冰系循环中，敌人身体、伤害数字、格挡/攻击文字和冰雾叠在一起，画面主次不稳定。",
        "把冰系技能拆成“短促接触爆点 + 敌人轮廓反馈 + 低透明度地面残留”。接触爆点可以亮，但持续雾气应降低 25%-40% 透明度，并给命中目标 0.12-0.2 秒描边或闪白。",
        boards["vfx"],
    )
    add_issue(
        doc,
        2,
        "敌人受击反应弱，打击缺少后果感",
        "多数敌人被命中后主要依赖伤害数字和血条表达结果，身体姿态、位移、硬直、局部反馈较弱。Boss 阶段尤其明显，体型巨大但被冰系命中时影响感不够。",
        "0:28-0:30 小蜘蛛死亡前的中间受击过程较轻；9:43-10:12 Boss 多次吃冰系攻击，但身体反馈与空间变化不足。",
        "建立命中等级：轻击短闪，小技能局部抖动，冰锥/重击触发硬直或局部冻结，Boss 不必被击退但要有头部/腹部/腿部局部受击、眼睛闪烁或甲壳震动。",
        boards["reaction"],
    )
    add_issue(
        doc,
        3,
        "敌人威胁可读性不足，受伤原因不够公平",
        "蛛网、毒液、红屏和控制效果已有结果表现，但攻击前摇、危险范围和来源提示不够稳定。玩家容易先看到自己掉血或被网住，再回想敌人做了什么。",
        "0:40-0:44、3:00-3:11、8:36-8:43 多次出现蛛网压屏或毒液效果，视觉结果强，但前置危险提示弱，且和玩家技能特效互相遮挡。",
        "蜘蛛吐网前增加 0.4-0.7 秒预兆：敌人抬腹/口器发光、屏幕边缘细网、地面或中心细线预警至少出现两项。玩家受击时加入来源方向、状态图标和短文字中的两项。",
        boards["threat"],
    )
    add_issue(
        doc,
        4,
        "法杖身份强，但长期占据命中阅读区",
        "第一人称法杖建立了角色身份，但法杖球体经常停在画面右中甚至中轴附近，压住敌人、血条、伤害数字或命中爆点。",
        "0:27-0:30、3:20-3:21、5:41、9:50-9:51 等时刻，法杖与敌人身体/血条重叠，影响玩家判断目标状态。",
        "施法前法杖可进入中心增强仪式感，释放后快速回到右下安全位；命中瞬间若敌人在中心，法杖可轻微让位或降低自身高光，保留身份但不抢主信息。",
        boards["staff"],
    )
    add_issue(
        doc,
        5,
        "多敌人战斗节奏和画面层级被重复冰系循环拉平",
        "中后段战斗大量由冰爆、冰锥、冻结、反弹、格挡和伤害数字组成。单次看有冲击力，但连续出现后，视觉峰值变成常态，难以区分普通命中、关键命中和高潮。",
        "6:34-7:18 的多敌人段落中，蓝白中心爆点、绿色回血/状态数字、红色攻击文字、格挡文字和敌人血条同时出现，信息密度持续偏高。",
        "为屏幕中央建立 VFX 预算：同一时刻最多一个主特效，其他持续层降级为背景残留。并为开场压制、敌人反扑、玩家爆发、收尾奖励设计不同强度曲线。",
        boards["rhythm"],
    )
    add_issue(
        doc,
        6,
        "Boss 阶段体型压迫强，但交互后果不够清晰",
        "红色蜘蛛 Boss 入场和体型压迫感成立，但战斗中 Boss 多次受击、喷吐、被冰冻时，玩家对“我是否打断/压制/削弱了它”的感知不足。",
        "9:40-10:13 Boss 多次吃冰系技能和出现绿色喷吐，但 Boss 身体变化较少，命中与威胁事件主要靠数字、红屏和特效表达。",
        "Boss 需要独立表现规则：大招前环境蛛网/卵囊变化，受击时局部受创，冻结时有局部冰壳沿身体扩散，破防或阶段转换时加入明显姿态和镜头节拍。",
        boards["boss"],
    )

    doc.add_heading("三、第一轮改动建议", level=1)
    doc.add_heading("快速收益", level=2)
    add_bullets(doc, [
        "所有敌人命中增加 0.12 秒左右的闪白或描边，Boss 改为局部闪白。",
        "冰雾、蛛网、毒液等持续遮挡特效透明度下调 25%-40%，接触爆点保留亮度但缩短生命周期。",
        "玩家受击反馈独立出来：红屏、来源方向、状态图标三选二，避免和敌人伤害数字混在一起。",
        "法杖释放后回到右下安全位，减少遮挡血条、命中点和敌人身体。",
        "蜘蛛网攻击增加 0.4-0.7 秒前摇，让玩家在被控前能读到危险。",
    ])
    doc.add_heading("较大改动", level=2)
    add_bullets(doc, [
        "建立命中类型规则：轻击、冰锥、冰爆、冻结、反弹、毒/网、Boss 命中、击杀分别有不同表现模板。",
        "建立 VFX 预算规则：屏幕中央只允许一个主特效，持续特效必须让出敌人剪影、血条和关键文字。",
        "建立敌人威胁规则：所有掉血/控制攻击都必须经历“预兆 -> 生效 -> 结果”三段。",
        "建立 Boss 交互规则：Boss 可不被击退，但必须通过局部反应、状态变化或节奏停顿表现“被影响”。",
    ])

    doc.add_heading("四、候选战斗表现框架规则", level=1)
    rules = [
        ("所有命中必须先让玩家看清“打中了谁”，再展示大范围特效。", "玩家技能、VFX、伤害数字", "静音缩小观看 3 敌人遭遇，仍能指出每次命中的目标。"),
        ("蜘蛛系控制技能必须有 0.4-0.7 秒可识别前摇。", "敌人、蛛网、毒液、玩家受击", "录制玩家被蛛网命中的片段，命中前应能判断“蜘蛛要放网了”。"),
        ("Boss 可以不被击退，但必须被影响。", "Boss 受击、冻结、重击、暴击", "录制 Boss 连续吃 3 次冰系技能，检查是否有局部受击、状态变化或节奏停顿。"),
        ("第一人称武器是身份层，不应长期占据命中阅读区。", "法杖、手部、镜头构图", "录制连续施法，确认法杖释放后回到安全区，不遮挡敌人血条和接触点。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, text in enumerate(["候选规则", "适用对象", "验证方式"]):
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_margins(hdr[i])
        hdr[i].paragraphs[0].add_run(text).bold = True
    for rule, applies, test in rules:
        row = table.add_row().cells
        for c in row:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c)
        row[0].text = rule
        row[1].text = applies
        row[2].text = test
    set_table_borders(table)

    doc.add_heading("五、下一轮录屏建议", level=1)
    add_numbered(doc, [
        "单只小蜘蛛 20-30 秒：只看普通攻击、受击、死亡与奖励反馈。",
        "三敌人混战 20-30 秒：专门检查 VFX 遮挡、伤害数字、目标辨识和节奏峰值。",
        "Boss 放网/毒液/被冰冻 20-30 秒：检查威胁前摇、玩家受击、Boss 受击后果。",
    ])
    note = doc.add_paragraph()
    note.add_run("备注：").bold = True
    note.add_run("本报告基于画面帧分析完成，未单独审听音频。SFX 的瞬态、材质层、混音优先级建议在下一轮配合有声录屏单独评审。")

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    BOARD_DIR.mkdir(parents=True, exist_ok=True)
    boards = {
        "vfx": make_board("01_vfx_hit_clarity", "问题 1：命中反馈被大面积 VFX 淹没", [
            (24.0, "冰系特效占据接触区"),
            (28.0, "敌人与文字叠在冰雾中"),
            (29.0, "命中数字可见但目标轮廓弱"),
            (397.0, "多敌人+冰雾+文字叠加"),
            (408.0, "格挡/攻击文字与特效重叠"),
            (436.0, "冰块遮挡敌人状态"),
        ]),
        "reaction": make_board("02_enemy_reaction", "问题 2：敌人受击反应弱，后果感不足", [
            (28.0, "小蜘蛛受击主要靠数字"),
            (29.0, "死亡前姿态变化较少"),
            (200.0, "大蜘蛛命中反馈偏轻"),
            (201.0, "击杀靠金币确认"),
            (583.0, "Boss 吃技能但身体反馈弱"),
            (596.0, "Boss 受击缺少局部后果"),
        ]),
        "threat": make_board("03_threat_readability", "问题 3：敌人威胁前摇与来源不够清晰", [
            (40.0, "毒液/危险结果出现"),
            (42.0, "蛛网压屏但前兆弱"),
            (182.0, "蛛网与敌人叠加"),
            (186.0, "控制结果强于预警"),
            (518.0, "Boss 网控覆盖屏幕"),
            (620.0, "Boss 阶段蛛网/受击混杂"),
        ]),
        "staff": make_board("04_staff_occlusion", "问题 4：法杖长期占据命中阅读区", [
            (27.0, "法杖压住中线目标"),
            (29.0, "法杖与伤害区重叠"),
            (200.0, "法杖遮挡敌人主体"),
            (341.0, "中心接触点被法杖抢视线"),
            (481.0, "多敌人时遮挡更明显"),
            (590.0, "Boss 血条/身体与法杖重叠"),
        ]),
        "rhythm": make_board("05_rhythm_vfx_budget", "问题 5：多敌人节奏与画面层级被拉平", [
            (394.0, "多敌人开场即高强度"),
            (402.0, "冰雾持续覆盖中心"),
            (410.0, "格挡/攻击/回血同时出现"),
            (418.0, "连续特效无明显峰谷"),
            (430.0, "手部蓄力与混战叠加"),
            (437.0, "冰块/文字/敌人同屏拥挤"),
        ]),
        "boss": make_board("06_boss_consequence", "问题 6：Boss 压迫强，但交互后果弱", [
            (580.0, "Boss 入场体型压迫成立"),
            (583.0, "首轮命中缺少局部反应"),
            (588.0, "大技能亮但 Boss 反馈弱"),
            (596.0, "Boss 被打中但状态变化不强"),
            (602.0, "玩家输出与 Boss 行为混叠"),
            (611.0, "未命中/受击信息靠文字表达"),
        ]),
    }
    path = build_doc(boards)
    print(path)
    for k, v in boards.items():
        print(k, v)


if __name__ == "__main__":
    main()
